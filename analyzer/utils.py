import os
import pdfplumber
from pdfminer.high_level import extract_text as pdfminer_extract_text
import docx

ALLOWED_EXTENSIONS = {'.pdf', '.docx'}
MAX_FILE_SIZE_MB = 5
MAX_FILE_SIZE_BYTES = MAX_FILE_SIZE_MB * 1024 * 1024


def validate_resume_file(uploaded_file):
    """
    Validates uploaded resume file extension and size.
    Returns (is_valid, error_message).
    """
    if not uploaded_file:
        return False, "No file uploaded. Please select a resume file."

    ext = os.path.splitext(uploaded_file.name)[1].lower()
    if ext not in ALLOWED_EXTENSIONS:
        return False, f"Invalid file format '{ext}'. Only PDF (.pdf) and Word (.docx) files are allowed."

    if uploaded_file.size > MAX_FILE_SIZE_BYTES:
        size_in_mb = round(uploaded_file.size / (1024 * 1024), 2)
        return False, f"File size ({size_in_mb} MB) exceeds the maximum allowed limit of {MAX_FILE_SIZE_MB} MB."

    return True, None


def extract_text_from_pdf(file_path):
    """
    Extracts plain text from a PDF file using pdfplumber with fallback to pdfminer.
    """
    extracted_text = ""

    # Primary method: pdfplumber
    try:
        with pdfplumber.open(file_path) as pdf:
            pages_text = []
            for page in pdf.pages:
                text = page.extract_text()
                if text:
                    pages_text.append(text.strip())
            extracted_text = "\n\n".join(pages_text).strip()
    except Exception as e:
        extracted_text = ""

    # Fallback method: pdfminer if pdfplumber extracted nothing or encountered an issue
    if not extracted_text:
        try:
            fallback_text = pdfminer_extract_text(file_path)
            if fallback_text and fallback_text.strip():
                extracted_text = fallback_text.strip()
        except Exception as e:
            pass

    if not extracted_text:
        raise ValueError("Unable to extract text from PDF. The file may be empty, image-only/scanned, or password protected.")

    return extracted_text


def extract_text_from_docx(file_path):
    """
    Extracts plain text from a DOCX file using python-docx.
    """
    try:
        doc = docx.Document(file_path)
        content_lines = []

        # Extract text from paragraphs
        for para in doc.paragraphs:
            clean_p = para.text.strip()
            if clean_p:
                content_lines.append(clean_p)

        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_cells:
                    content_lines.append(" | ".join(row_cells))

        extracted_text = "\n".join(content_lines).strip()
        if not extracted_text:
            raise ValueError("The uploaded DOCX file appears to be empty or contains no readable text.")

        return extracted_text
    except Exception as e:
        if isinstance(e, ValueError):
            raise e
        raise ValueError(f"Failed to read DOCX file: {str(e)}")


def extract_resume_text(file_path, file_extension):
    """
    Main helper function to extract text based on file extension.
    """
    ext = file_extension.lower()
    if ext == '.pdf':
        return extract_text_from_pdf(file_path)
    elif ext == '.docx':
        return extract_text_from_docx(file_path)
    else:
        raise ValueError(f"Unsupported file format: {ext}")
